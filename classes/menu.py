from datetime import datetime, date
import re
from typing import List, Optional
from docx import Document  # type: ignore
from classes.recipe import Recipe


class Menu:
    def __init__(self, menu_file: str) -> None:
        self.menu_file: str = menu_file
        self.recipes: List[Recipe] = []
        self.menu_date_string: Optional[str] = None
        self.menu_date: Optional[date] = None
        self.get_menu_date()

    def get_menu_date(self) -> None:
        match = re.search(r"\s(\d{6})\s", self.menu_file)
        if match:
            menu_date_str: str = match.group(1)  # type explicitly as str
            self.menu_date_string = menu_date_str
            self.menu_date = datetime.strptime(menu_date_str, "%y%m%d").date()
        else:
            self.menu_date_string = None
            self.menu_date = None

    def analyse(self) -> None:
        """
        Analyse the Word document, extract recipes from the first table,
        and store Recipe objects in self.recipes.
        """
        doc: Document = Document(self.menu_file)  # type: ignore
        if not doc.tables:
            raise ValueError("No tables found in the document.")

        table = doc.tables[0]
        cells: List[str] = [str(cell.text).strip() for row in table.rows for cell in row.cells]
        cells = cells[:7]

        for cell_text in cells:
            recipe: Recipe = Recipe(cell_text, self.menu_date)
            recipe.parse()
            if recipe.recipe not in ["Tea", "Tea:"]:
                self.recipes.append(recipe)
