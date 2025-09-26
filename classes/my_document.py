from classes.environment_variable import EnvironmentVariable
from classes.qr_maker import QrMaker
from docx import Document
from docx.shared import Cm
import os
import sys


class MyDocument(object):
    def __init__(self, date_string):
        self.date_string = date_string
        self.get_folders_and_filenames()

    def get_folders_and_filenames(self):
        self.year = "20" + self.date_string[0:2]
        shopping_folder = str(EnvironmentVariable("shopping_folder", "string", False).value)
        output_filename_template = str(EnvironmentVariable("output_filename_template", "string", False).value)
        output_filename = output_filename_template.format(date_string=self.date_string)
        self.filename = os.path.join(shopping_folder, self.year, output_filename)
        if not (os.path.exists(self.filename)):
            print("File cannot be found, would you like to generate a blank file instead?")
            sys.exit()

    def cell_contains_hyperlink(self, cell, part):
        """
        Check if a cell contains a hyperlink.
        Returns (True, [urls]) if found, else (False, []).
        """
        el = cell._element
        urls = []

        # 1) Explicit <w:hyperlink r:id="...">
        for hl in el.xpath('.//w:hyperlink'):
            rId = hl.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
            if rId and rId in part.rels:
                urls.append(part.rels[rId].target_ref)

        # 2) Field-based hyperlinks
        for fld in el.xpath('.//w:fldSimple'):
            instr = fld.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}instr')
            if instr and 'HYPERLINK' in instr.upper():
                parts = instr.split('"')
                if len(parts) >= 2:
                    urls.append(parts[1])

        # 3) Instruction text pieces
        for it in el.xpath('.//w:instrText'):
            if it.text and 'HYPERLINK' in it.text.upper():
                parts = it.text.split('"')
                if len(parts) >= 2:
                    urls.append(parts[1])

        # 4) Plain text URLs
        text = (cell.text or "").strip()
        if any(prefix in text.lower() for prefix in ["http://", "https://", "www."]):
            urls.append(text)

        return (len(urls) > 0, urls)

    def mark_links_with_image(self):
        doc = Document(self.filename)
        if not doc.tables:
            print("No tables found in document.")
            return

        table = doc.tables[0]
        part = doc.part  # needed for hyperlink relationships

        to_mark = []  # list of (row_index, col_index, urls)

        # Scan all cells
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                has_link, urls = self.cell_contains_hyperlink(cell, part)
                if len(urls) > 0:
                    qr_maker = QrMaker(urls[0])
                    qr_maker.make_qr_code()

                    if has_link:
                        print(f"Row {r_idx}, Col {c_idx} contains hyperlink(s): {urls}")
                        to_mark.append((r_idx, c_idx, urls, qr_maker.filename))

        if not to_mark:
            print("No hyperlinks found in first table.")
            return

        # Insert image in the left cell
        for r_idx, c_idx, urls, filename in to_mark:
            if c_idx == 0:
                continue
            target_row = table.rows[r_idx]
            if c_idx - 1 < len(target_row.cells):
                left_cell = target_row.cells[c_idx - 1]
                left_cell.text = ""  # clear existing text
                run = left_cell.paragraphs[0].add_run()
                run.add_picture(filename, width=Cm(1.25), height=Cm(1.25))

        overwrite_original = EnvironmentVariable("overwrite_original", "boolean", True).value
        if overwrite_original:
            output_path = self.filename
        else:
            output_path = self.filename.replace(".docx", "-qr.docx")
        doc.save(output_path)

        print(f"Saved modified document as: {output_path}")
