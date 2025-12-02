from classes.environment_variable import EnvironmentVariable
from datetime import datetime, timedelta
from docx import Document
import os
import sys
import utils.utils as u
import subprocess


# Ensure correct command-line usage
if len(sys.argv) < 2:
    # print("Usage: python script.py YY-MM-DD {no-qr}")
    # sys.exit(1)
    date_arg = u.next_saturday()
else:
    # Parse the input date
    date_arg = str(sys.argv[1])

try:
    input_date = datetime.strptime(date_arg, "%y-%m-%d")
    has_hyphens = True
except ValueError:
    try:
        input_date = datetime.strptime(date_arg, "%y%m%d")
        has_hyphens = False
    except ValueError:
        print("Error: Invalid date format. Use YY-MM-DD.")
        sys.exit(1)


# Get the document save path
shopping_folder = EnvironmentVariable("shopping_folder", "string", False).value
year_folder = "20" + date_arg[0:2]
output_path = os.path.join(str(shopping_folder), year_folder)

# Get the document save filename
date_string = date_arg.replace("-", "")
output_filename_template = str(EnvironmentVariable("output_filename_template", "string", False).value)
output_filename = output_filename_template.format(date_string=date_string)

# Get the full path
output_path = os.path.join(output_path, output_filename)

if os.path.exists(output_path):
    print("\n\nThe file already exists. Overwrite? (Y/N)")
    choice = u.getch().lower()
    print(choice)  # echo the key the user pressed
    if choice != "y":
        print("\nAborted.")
        sys.exit()
    print("\nProceeding...")

# Parse the no-qr flag
generate_qrs = True
if len(sys.argv) >= 3:
    no_qr = sys.argv[2]
    if no_qr in ["no_qr", "no-qr"]:
        generate_qrs = False

# Function to format the date as 'Saturday, 1st Jan 2025'


def format_date(dt: datetime) -> str:
    day = dt.day
    suffix = "th" if 11 <= day <= 13 else {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
    formatted_day = dt.strftime("%a").upper()  # Day in uppercase
    # Suffix and month in original case
    return f"{formatted_day} {day}{suffix} {dt.strftime('%b')}"


# Load the Word document from template
if generate_qrs:
    template_filename = "template-qr.docx"
else:
    template_filename = "template.docx"
template_path = os.path.join(os.getcwd(), "templates", template_filename)
doc = Document(template_path)

# Ensure the document has at least one table
if not doc.tables:
    print("Error: No tables found in the template document.")
    sys.exit(1)

table = doc.tables[0]  # Get the first table

# Insert dates and "Lunch" into the table cells (row 1, 2, 3, 4 - 2 cells per row)
for row in range(4):
    for column in range(2):
        if generate_qrs:
            actual_column_index = (column * 2) + 1
        else:
            actual_column_index = column
        cell_index = row + (column * 4)
        formatted_date = format_date(input_date + timedelta(days=cell_index))
        cell = table.rows[row].cells[actual_column_index]

        if row * column < 3:
            # Clear existing content
            u.clear_cell(cell)

            # Add formatted date with "Days" style
            date_paragraph = cell.add_paragraph(formatted_date)
            date_paragraph.style = "Days"

            # Add "Lunch" and tea below the date with "Meals" style
            lunch_paragraph = cell.add_paragraph("Lunch:\t")
            lunch_paragraph.style = "Meals"
            tea_paragraph = cell.add_paragraph("Tea:\t")
            tea_paragraph.style = "Meals"

# Get the document save path
shopping_folder = EnvironmentVariable("shopping_folder", "string", False).value
year_folder = "20" + date_arg[0:2]
output_path = os.path.join(str(shopping_folder), year_folder)

# Get the document save filename
date_string = date_arg.replace("-", "")
output_filename_template = str(EnvironmentVariable("output_filename_template", "string", False).value)
output_filename = output_filename_template.format(date_string=date_string)

# Get the full path
output_full_path = os.path.join(output_path, output_filename)

# Save the document
doc.save(output_full_path)
print(f"Document saved as {output_full_path}")
subprocess.run(["open", output_path])
